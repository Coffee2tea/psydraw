import base64
import os
import shutil
import tempfile
import time
import zipfile
from io import BytesIO

import streamlit as st
from docx import Document
from PIL import Image

# Use our custom ChatOpenAI wrapper instead of the original
import sys

# Add parent directory to path for imports
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../..")))

from src.custom_chat_openai import ChatOpenAI
from src.model_langchain import HTPModel

# Add monkey patch to disable proxies in OpenAI
import openai
# Override any proxy settings that might be configured in the environment or elsewhere
if hasattr(openai, '_client'):
    if hasattr(openai._client, 'proxies'):
        delattr(openai._client, 'proxies')

def img_to_bytes(img_path):
    img = Image.open(img_path)
    buffered = BytesIO()
    img.save(buffered, format=img.format)
    return base64.b64encode(buffered.getvalue()).decode()

# Setup paths
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(os.path.dirname(SCRIPT_DIR))
ASSETS_DIR = os.path.join(PROJECT_ROOT, "assets")

def get_asset_path(filename):
    return os.path.join(ASSETS_DIR, filename)

SUPPORTED_LANGUAGES = {
    "English": "en"
}

LANGUAGES = {
    "en": {
        "analysis_settings": "Analysis Settings",
        "model_settings": "🍓 Model Settings",
        "batch_title": "📊 Batch Analysis",
        "language_label": "Language:",
        "select_folder": "Enter the folder path containing images:",
        "no_images_found": "No image files found in the selected folder.",
        "images_found": "{} image files found. Ready for batch analysis.",
        "start_batch_analysis": "Start Batch Analysis",
        "batch_results_summary": "Batch Analysis Results Summary",
        "download_batch_results": "Download Batch Results",
        "enter_valid_folder": "Please upload images.",
        "error_no_api_key": "❌ Please enter your API key in the sidebar before starting the analysis.",
        "batch_instructions_title": "📋 Batch Analysis Instructions",
        "upload_images": "Upload Images for Batch Analysis",
        "images_uploaded": "{} images uploaded successfully.",
        "upload_images_prompt": "Please upload images to start batch analysis.",
    "batch_instructions": """
    **Please read the following instructions carefully before proceeding with batch analysis:**

    1. **API Key**: Ensure you have filled in your API key in the sidebar. This is crucial for the analysis to work.
    
    2. **Preparation**: 
       - Prepare the images you want to analyze on your local device.
       - Make sure all images are in .jpg, .jpeg, or .png format.
    
    3. **Image Upload**: 
       - Click on the 'Upload Images' button or drag and drop your images into the designated area.
       - You can select multiple images at once for batch processing.
    
    4. **Time Consideration**: Batch analysis may take a considerable amount of time, depending on the number and size of images. Please be patient.
    
    5. **Network and API Credits**:
       - Ensure you have a stable internet connection throughout the process.
       - Check that you have sufficient API credits for the entire batch. Each image consumes credits.
    
    6. **Starting Analysis**: 
       - After uploading your images, click on the 'Start Batch Analysis' button to begin the process.
    
    7. **Results**: 
       - Once the analysis is complete, use the 'Download Batch Results' button to save the full analysis results as a zip file.
       - The zip file will contain individual reports for each image and a summary of any failed analyses.

    **Note**: This tool is for reference only and cannot replace professional psychological evaluation. If you have concerns, please consult a qualified mental health professional.
    """,
    "welcome": "Welcome to the Batch Analysis Page",
    "batch_results": "Batch Analysis Finished, Please download the results. Successful: {} | Failed: {}",
    "download_batch_results": "Download Batch Results (ZIP)",
    "ai_disclaimer": "NOTE: AI-generated content, for reference only. Not a substitute for medical diagnosis.",
    }
}

@st.cache_resource
def get_uploaded_files():
    return []

def pil_to_base64(image: Image.Image, format: str = "JPEG") -> str:
    """Convert PIL image to base64 string."""
    buffered = BytesIO()
    image.save(buffered, format=format)
    return base64.b64encode(buffered.getvalue()).decode('utf-8')

def get_text(key):
    return LANGUAGES[st.session_state['language_code']][key]

def save_results(results):
    with tempfile.TemporaryDirectory() as temp_dir:
        for result in results:
            file_name = result['file_name']
            file_name_without_ext = os.path.splitext(file_name)[0]
            
            result_folder = os.path.join(temp_dir, file_name_without_ext)
            os.makedirs(result_folder, exist_ok=True)
            
            # 保存上传的图片
            if result['image']:
                image_path = os.path.join(result_folder, file_name)
                result['image'].save(image_path)
            
            doc = Document()
            if result['success']:
                doc.add_paragraph(get_text("ai_disclaimer"))
                if result['analysis_result']['classification'] is True:
                    signal = result['analysis_result']['signal']
                    final = result['analysis_result']['final']
                    doc.add_paragraph(signal)
                    doc.add_paragraph(final)
                else:
                    signal = result['analysis_result']['fix_signal']
                    doc.add_paragraph(signal)
            else:
                doc.add_paragraph("failed")
            doc_path = os.path.join(result_folder, f"{file_name_without_ext}.docx")
            doc.save(doc_path)
        
        failed_path = os.path.join(temp_dir, "failed.txt")
        with open(failed_path, "w") as f:
            for result in results:
                if not result['success']:
                    f.write(f"{result['file_name']}\n")
                    
        zip_path = os.path.join(temp_dir, "results.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for root, _, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, arcname)
        with open(zip_path, "rb") as f:
            zip_content = f.read()
    
    return zip_content    
        
def batch_analyze(uploaded_files):
    results = []
    
    MULTIMODAL_MODEL = "gpt-4-vision-preview"
    TEXT_MODEL = "gpt-4-turbo"
    
    # Check if API key is set in environment
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        st.error("❌ Please enter your API key in the sidebar before starting the analysis.")
        return [], 0
    
    # Create custom ChatOpenAI instances with explicit parameters
    try:
        print(f"Initializing ChatOpenAI with API Key: {'Yes (length: ' + str(len(api_key)) + ')' if api_key else 'No'}")
        print(f"Base URL: {st.session_state.base_url or 'default'}")
        
        text_model = ChatOpenAI(
            model_name=TEXT_MODEL,
            api_key=api_key,
            base_url=st.session_state.base_url,
            temperature=0.2,
        )
        
        multimodal_model = ChatOpenAI(
            model_name=MULTIMODAL_MODEL,
            api_key=api_key,
            base_url=st.session_state.base_url,
            temperature=0.2,
        )
        
        model = HTPModel(
            text_model=text_model,
            multimodal_model=multimodal_model,
            language=st.session_state['language_code'],
            use_cache=True
        )
        print("HTPModel initialized successfully")
    except Exception as e:
        print(f"Error during initialization: {str(e)}")
        st.error(f"Error initializing models: {str(e)}")
        return [], 0
        
    progress_bar = st.progress(0, text=f"Progressing: 0/{len(uploaded_files)}")
    start_time = time.time()
    success = 0
    for i, uploaded_file in enumerate(uploaded_files):
        try:
            print(f"Processing file {i+1}/{len(uploaded_files)}: {uploaded_file.name}")
            image = Image.open(uploaded_file)
            image_data = pil_to_base64(image)
            print(f"Image converted to base64 string (length: {len(image_data)})")
            
            print(f"Starting workflow analysis with language: {st.session_state['language_code']}")
            response = model.workflow(image_path=image_data, language=st.session_state['language_code'])
            print(f"Analysis completed successfully for {uploaded_file.name}")
            
            results.append({
                "file_name": uploaded_file.name,
                "analysis_result": response,
                "success": True,
                "image": image
            })
            success += 1
        except Exception as e:
            print(f"Error processing {uploaded_file.name}: {str(e)}")
            import traceback
            print(f"Traceback: {traceback.format_exc()}")
            
            results.append({
                "file_name": uploaded_file.name,
                "analysis_result": str(e),
                "success": False,
                "image": image
            })
        
        elapsed_time = time.time() - start_time
        progress = (i + 1) / len(uploaded_files)
        estimated_total_time = elapsed_time / progress if progress > 0 else 0
        remaining_time = estimated_total_time - elapsed_time
        
        elapsed_str = time.strftime("%H:%M:%S", time.gmtime(elapsed_time))
        remaining_str = time.strftime("%H:%M:%S", time.gmtime(remaining_time))
        
        progress_bar.progress(progress, text=f"Progressing: {i + 1}/{len(uploaded_files)} | Elapsed: {elapsed_str} | Remaining: {remaining_str}")
    
    st.success(get_text("batch_results").format(success, len(uploaded_files) - success))
    
    return results, success

def sidebar() -> None:
    """Render sidebar components."""
    with st.sidebar:
        logo_path = get_asset_path("logo-3.png")
        st.markdown(
            f'<img src="data:image/png;base64,{img_to_bytes(logo_path)}" alt="logo" width="100%">',
            unsafe_allow_html=True
        )
        
    st.sidebar.markdown(f"## {get_text('model_settings')}")
    
    # API Key input
    api_key = st.sidebar.text_input(
        "OpenAI API Key:",
        value=os.getenv("OPENAI_API_KEY", ""),
        type="password",
        help="Enter your OpenAI API key here. It will only be stored for the duration of this session."
    )
    
    if api_key:
        os.environ["OPENAI_API_KEY"] = api_key
        st.session_state.api_key = api_key
    
    # Base URL input (optional)
    base_url = st.sidebar.text_input(
        "OpenAI API Base URL (Optional):",
        value=os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1"),
        help="Custom API endpoint. Leave as default unless you're using a proxy."
    )
    
    if base_url:
        os.environ["OPENAI_BASE_URL"] = base_url
        st.session_state.base_url = base_url
    
    # Always set language to English
    st.session_state['language'] = "English"
    st.session_state['language_code'] = "en"
    
    st.sidebar.markdown("---")
    if st.sidebar.button("Start Analysis"):
        st.session_state.start_analysis = True
    
def batch_page():
    st.title(get_text("batch_title"))
    
    st.write(get_text("welcome"))
    
    with st.expander(get_text("batch_instructions_title"), expanded=True):
        st.markdown(get_text("batch_instructions"))
    
    
def main():
    st.set_page_config(page_title="PsyDraw: Batch Analysis", page_icon="📊", layout="wide")
    
    if 'language_code' not in st.session_state:
        st.session_state['language_code'] = "en"

    if 'language' not in st.session_state:
        st.session_state['language'] = "English"
        
    # Initialize language selector
    if 'language_selector' not in st.session_state:
        st.session_state['language_selector'] = st.session_state['language']
    
    sidebar()
    batch_page()
    
    uploaded_files = st.file_uploader(get_text("upload_images"), accept_multiple_files=True, type=['png', 'jpg', 'jpeg'], key="file_uploader")
    status_placeholder = st.empty()
    if uploaded_files:
        cached_files = get_uploaded_files()
        cached_files.extend(uploaded_files)
        status_placeholder.success(get_text("images_uploaded").format(len(cached_files)))
        
    if st.session_state.get('start_analysis'):
        # Check if API key is set in environment
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            st.error("❌ Please enter your API key in the sidebar before starting the analysis.")
        elif uploaded_files:
            results, success = batch_analyze(uploaded_files=uploaded_files)
            
            zip_content = save_results(results)
            st.download_button(
                label = get_text("download_batch_results"),
                data=zip_content,
                file_name="batch_analysis_results.zip",
                mime="application/zip"
            )
        st.session_state.start_analysis = False
    
    

if __name__ == "__main__":
    main()